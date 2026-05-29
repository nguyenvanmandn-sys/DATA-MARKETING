# -*- coding: utf-8 -*-
"""Tao file Word: Nghien cuu thi truong iPhone cu & dinh huong content."""
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = Document()

# ---- Style co ban ----
style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(11)

NAVY = RGBColor(0x1F, 0x3A, 0x5F)
RED = RGBColor(0xC0, 0x2B, 0x2B)
GREEN = RGBColor(0x1E, 0x7A, 0x3C)

def heading(text, level=1, color=NAVY):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = color
    return h

def bullet(text, bold_prefix=None):
    p = doc.add_paragraph(style='List Bullet')
    if bold_prefix:
        r = p.add_run(bold_prefix)
        r.bold = True
        p.add_run(text)
    else:
        p.add_run(text)
    return p

def para(text, bold=False, italic=False, color=None, align=None):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = bold
    r.italic = italic
    if color:
        r.font.color.rgb = color
    if align:
        p.alignment = align
    return p

# ====================== TRANG BIA ======================
t = doc.add_paragraph()
t.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = t.add_run('NGHIEN CUU THI TRUONG & DINH HUONG CONTENT')
r.bold = True
r.font.size = Pt(20)
r.font.color.rgb = NAVY

t2 = doc.add_paragraph()
t2.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = t2.add_run('Nganh: Kinh doanh iPhone cu (used iPhone)')
r.font.size = Pt(14)
r.italic = True

t3 = doc.add_paragraph()
t3.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = t3.add_run('Khu vuc tham chieu: Da Nang  |  Phuc vu du an DATA-MARKETING')
r.font.size = Pt(12)

doc.add_paragraph()
para('Ngay lap: 29/05/2026', italic=True, align=WD_ALIGN_PARAGRAPH.CENTER)

# Ghi chu pham vi / phuong phap
doc.add_paragraph()
note = doc.add_paragraph()
rn = note.add_run('LUU Y VE PHUONG PHAP & PHAM VI')
rn.bold = True
rn.font.color.rgb = RED
para('Bao cao nay KHONG su dung du lieu cao (scrape) tu trang Facebook cua bat ky shop cu the nao. '
     'Facebook chan thu thap tu dong va viec cao binh luan/danh gia cua mot trang ben thu ba vi pham '
     'Dieu khoan dich vu cua Meta, dong thoi cham toi du lieu ca nhan cua nguoi binh luan. '
     'Thay vao do, bao cao tong hop tu cac nguon bao chi va trang chuyen mon cong khai ve nganh iPhone cu, '
     'nham giup ban xay dung NOI DUNG / CONTENT mot cach hop le.', italic=True)
para('Cach dung dung dan: dung phan "Nhung diem khach phan nan" de hieu noi dau cua khach, roi xay '
     'content cho thay SHOP CUA BAN giai quyet duoc nhung noi dau do. Khong dung de cong kich doi thu.',
     italic=True, color=GREEN)

doc.add_page_break()

# ====================== 1. TONG QUAN ======================
heading('1. Tong quan thi truong iPhone cu', 1)
para('iPhone cu la phan khuc nong vi gia mem hon may moi nhung van giu gia tri thuong hieu Apple. '
     'Tuy nhien day cung la phan khuc nhieu rui ro lua dao nhat, nen NIEM TIN la yeu to quyet dinh chot don. '
     'Nguoi mua thuong lo lang truoc 4 nhom van de: chat luong may (hang dung), pin, bao hanh/doi tra, va uy tin cua nguoi ban.')

# ====================== 2. KHACH QUAN TAM ======================
heading('2. Nhung dieu nguoi mua iPhone cu quan tam nhat', 1)

heading('2.1. Chat luong & nguon goc may', 2)
bullet('lieu may co phai hang ZIN (nguyen ban) hay hang dung, da thay vo / linh kien.', 'Hang ZIN hay hang dung: ')
bullet('may con dinh iCloud thi khong the su dung - day la noi so lon nhat.', 'iCloud: ')
bullet('khach muon biet may quoc te (LL/A) hay Lock, vi anh huong song & gia tri.', 'May quoc te hay Lock: ')

heading('2.2. Tinh trang pin', 2)
bullet('khach quan tam % pin (battery health). Duoi 85% bi coi la yeu.', 'Do chai pin: ')
bullet('lo so mua phai may da "kich pin" gia 100% nhung thuc te pin da chai.', 'Pin that hay pin do: ')

heading('2.3. Gia ca', 2)
bullet('khach luon so sanh gia cung mau may o nhieu noi; gia qua re lai gay nghi ngo hang dom.', 'So sanh mat bang gia: ')
bullet('tra gop, thu cu doi moi la diem cong lon.', 'Chinh sach gia: ')

heading('2.4. Bao hanh & doi tra', 2)
bullet('thoi gian (6-12 thang), bao hanh nhung gi, co bao hanh pin khong.', 'Chinh sach bao hanh ro rang: ')
bullet('doi tra trong 7-30 ngay neu loi giup khach yen tam xuong tien.', 'Doi tra: ')

heading('2.5. Uy tin nguoi ban', 2)
bullet('cua hang co dia chi that, co the den tan noi kiem tra may.', 'Co dia chi vat ly: ')
bullet('phan hoi that tu khach cu, hinh anh / video thuc te.', 'Bang chung xa hoi: ')

# ====================== 3. PHAN NAN ======================
heading('3. Nhung diem khach hang thuong PHAN NAN voi shop iPhone cu', 1)
para('Day la cac noi dau pho bien (rut ra tu canh bao bao chi & cong dong), khong gan voi mot shop cu the. '
     'Hieu duoc chung giup ban viet content "danh trung" lo lang cua khach:', italic=True)

bullet('Quang cao "ZIN 100%" nhung thuc te la may dung, thay man / thay vo.', 'Hang dung doi lot ZIN: ')
bullet('Reset xong may doi Apple ID cu - khach phat hien sau khi da tra tien.', 'Dinh iCloud: ')
bullet('Hien thi pin 100% nhung dung nhanh het - chieu "kich pin".', 'Pin gia: ')
bullet('Bao hanh mo ho, lan loi sang loi khac, hoac vien co tinh tu choi.', 'Bao hanh khong dung cam ket: ')
bullet('Quang cao gia hoi roi thu them phu phi, hoac gia tren bai khong dung gia ban.', 'Gia ao / phu phi an: ')
bullet('Giao hang qua mang, bat tra tien truoc roi giao may khac / may hong (lua dao ship COD).', 'Ship roi giao may khac: ')
bullet('Nhan tin cham, thai do thieu nhiet tinh, khong tu van ky truoc khi mua.', 'Cham phan hoi / thai do: ')
bullet('Sau khi ban xong thi "lo" khach, kho lien lac khi can bao hanh.', 'Dich vu hau mai kem: ')

# ====================== 4. DINH HUONG CONTENT ======================
heading('4. Dinh huong CONTENT cho shop (bien noi dau thanh diem manh)', 1)

heading('4.1. Tru cot thong diep', 2)
bullet('Moi may deu co video / hinh test thuc te (3uTools, tinh trang pin, IMEI).', 'Minh bach: ')
bullet('Cam ket ZIN - sai den 1 chi tiet hoan tien 100%%.', 'Cam ket manh: ')
bullet('Bao hanh ro rang (vd 12 thang), doi tra 7-30 ngay, ho tro tra gop.', 'Hau mai: ')
bullet('Dia chi cua hang ro rang tai Da Nang - moi khach den kiem tra truc tiep.', 'Hien dien that: ')

heading('4.2. Cac dang bai nen dang (content pillar)', 2)
bullet('"5 cach kiem tra iPhone cu truoc khi mua" - dang chuyen gia, tao niem tin.', 'Bai kien thuc: ')
bullet('Quay clip test pin, iCloud, camera tung may dang ban.', 'Bai review may that: ')
bullet('Chup tin nhan / danh gia that cua khach da mua (xin phep truoc).', 'Bai feedback khach: ')
bullet('Ke cau chuyen khach tung mua phai hang dung noi khac, ve voi shop hai long.', 'Bai cau chuyen: ')
bullet('Bang gia cap nhat, chuong trinh thu cu doi moi, tra gop 0%%.', 'Bai uu dai: ')

heading('4.3. Cong thuc bai ban hang (AIDA / PAS)', 2)
para('PAS - Problem / Agitate / Solve:', bold=True)
bullet('Goi dung noi so: "Mua iPhone cu so nhat dieu gi? Hang dung va pin chai phai khong?"')
bullet('Khoet sau: "Rat nhieu nguoi mat tien trieu vi may dung doi lot ZIN..."')
bullet('Giai phap: "Tai [TEN SHOP], moi may deu test live truoc mat ban, sai hoan tien 100%%."')
para('Ket bai bang CTA ro rang: "Inbox ngay de nhan bang gia hom nay" + tao cam giac co han (so luong / uu dai).',
     italic=True)

heading('4.4. Nhung tu khoa / cum tu nen dung', 2)
bullet('ZIN nguyen ban, may quoc te, bao hanh 12 thang, doi tra 1-1, test truc tiep, '
       'tra gop 0%%, thu cu len doi, cam ket hoan tien.')

# ====================== 5. CHECKLIST ======================
heading('5. Checklist trien khai content tuan dau', 1)
for item in [
    'Viet 1 bai "cam nang kiem tra iPhone cu" ghim len dau trang.',
    'Quay 3-5 clip test may that dang ban trong tuan.',
    'Thu thap & xin phep dang 5 feedback khach cu.',
    'Lam 1 anh bang gia cap nhat + chinh sach bao hanh ro rang.',
    'Soan 3 mau tin nhan tu van nhanh (chao - tu van - chot).',
    'Len lich dang: xen ke kien thuc / review / feedback / uu dai (tranh chi ban hang).',
]:
    p = doc.add_paragraph(style='List Bullet')
    p.add_run('[ ] ').bold = True
    p.add_run(item)

# ====================== NGUON ======================
heading('Nguon tham khao (cong khai)', 1)
sources = [
    'Cach mua iPhone cu khong bi lua - 24hStore: https://24hstore.vn/huong-dan-ky-thuat/cach-mua-iphone-cu-khong-bi-lua-n8225',
    '8 luu y khi mua iPhone cu - CellphoneS/Sforum: https://cellphones.com.vn/sforum/luu-y-khi-mua-iphone-cu',
    'Cach tranh bi lua khi mua iPhone da qua su dung - Thanh Nien: https://thanhnien.vn/cach-tranh-bi-lua-khi-mua-iphone-da-qua-su-dung-185231122114859418.htm',
    'Kich pin - chieu tro nguoi mua iPhone cu can canh giac - XTmobile: https://www.xtmobile.vn/kich-pin-chieu-tro-nguoi-mua-iphone-cu-can-phai-canh-giac-de-tranh-bi-lua',
    'Chieu lua pin 100%% tren iPhone cu - vietnam.vn: https://www.vietnam.vn/en/chieu-lua-pin-100-tren-iphone-cu-khien-nguoi-mua-sap-bay',
    'Canh giac lua dao mua iPhone gia re - Cong an Nghe An: https://congan.nghean.gov.vn/thong-tin-chuyen-de/canh-bao-toi-pham/202503/canh-giac-lua-dao-mua-iphone-chinh-hang-gia-re-1037063/',
    'iPhone tra bao hanh la gi - Thegioididong: https://www.thegioididong.com/hoi-dap/iphone-tra-bao-hanh-la-gi-1580275',
    'Content ban hang Facebook 15 cong thuc - Adsplus: https://adsplus.vn/blog/content-ban-hang-facebook-15-cong-thuc-viet-bai-chot-don-hieu-qua',
    '20+ mau content ban dien thoai - AgencyVN: https://agencyvn.com/content-ban-dien-thoai',
]
for s in sources:
    p = doc.add_paragraph(style='List Bullet')
    p.add_run(s).font.size = Pt(10)

out = '/home/user/DATA-MARKETING/Nghien_cuu_iPhone_cu_va_Content.docx'
doc.save(out)
print('Saved:', out)
