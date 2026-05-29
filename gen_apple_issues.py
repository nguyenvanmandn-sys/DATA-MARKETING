# -*- coding: utf-8 -*-
"""Bao cao: Cac van de nguoi dung iPhone/iPad/MacBook dau dau & ban tan tren MXH."""
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = Document()
style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(11)

NAVY = RGBColor(0x1F, 0x3A, 0x5F)
RED = RGBColor(0xC0, 0x2B, 0x2B)
ORANGE = RGBColor(0xD1, 0x7A, 0x00)
GREEN = RGBColor(0x1E, 0x7A, 0x3C)

def heading(text, level=1, color=NAVY):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = color
    return h

def bullet(text, bold_prefix=None):
    p = doc.add_paragraph(style='List Bullet')
    if bold_prefix:
        r = p.add_run(bold_prefix); r.bold = True
        p.add_run(text)
    else:
        p.add_run(text)
    return p

def para(text, bold=False, italic=False, color=None, align=None):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = bold; r.italic = italic
    if color: r.font.color.rgb = color
    if align: p.alignment = align
    return p

def hot(level):
    # tra ve nhan muc do nong
    m = {'cao': ('[NONG - rat nhieu bai/binh luan]', RED),
         'tb': ('[TRUNG BINH]', ORANGE),
         'thap': ('[THUONG XUYEN]', GREEN)}
    txt, col = m[level]
    p = doc.add_paragraph()
    r = p.add_run(txt); r.bold = True; r.font.size = Pt(9); r.font.color.rgb = col
    return p

# ===== BIA =====
t = doc.add_paragraph(); t.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = t.add_run('CAC VAN DE NGUOI DUNG IPHONE / IPAD / MACBOOK'); r.bold = True; r.font.size = Pt(19); r.font.color.rgb = NAVY
t2 = doc.add_paragraph(); t2.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = t2.add_run('Quet tu cac chu de dang noi & duoc ban tan tren MXH (Facebook / TikTok / YouTube / Reddit / dien dan)'); r.italic = True; r.font.size = Pt(12)
doc.add_paragraph()
para('Ngay lap: 29/05/2026  |  Du an DATA-MARKETING', align=WD_ALIGN_PARAGRAPH.CENTER, italic=True)

doc.add_paragraph()
nt = doc.add_paragraph(); rn = nt.add_run('LUU Y VE PHUONG PHAP'); rn.bold = True; rn.font.color.rgb = RED
para('Khong cao truc tiep binh luan tu TikTok / Facebook / YouTube (cac nen tang chan thu thap tu dong va viec do '
     'vi pham Dieu khoan dich vu). Bao cao tong hop tu bao chi cong nghe & trang chuyen mon cong khai, phan anh dung '
     'nhung van de dang duoc nguoi dung ban tan nhieu nhat tren MXH. Nhan [NONG] = chu de co rat nhieu bai/binh luan.',
     italic=True)
doc.add_page_break()

# ===== TOM TAT TOP =====
heading('0. TOP van de "nhuc nhoi" nhat (tong hop nhanh)', 1)
for x in [
    ('Hao pin & nong may sau cap nhat iOS', 'Apple da THUA NHAN iOS 26.1 gay hao pin, nong may. Chu de bi than phien nhieu nhat moi dot update.'),
    ('iPhone 17 de tray xuoc ("Scratchgate")', 'Vien nhom anode de troc/xuoc, nhat la ban mau dam (Pro xanh dam, Air den). Hashtag >40 trieu luot xem tren Weibo.'),
    ('Wi-Fi / Bluetooth chap chon tren iPhone 17', 'Song yeu hon doi truoc, Wi-Fi tu ngat-ket noi lai; hang tram binh luan tren MacRumors, Reddit, Apple Community.'),
    ('Pin MacBook bi phong', 'Ban phim/nap lung troi len, nguy hiem, chi phi thay cao.'),
    ('iPad/iPhone treo tao, sac khong vao', 'Loi pho bien sau thoi gian dung; gay mat du lieu, lo lang.'),
]:
    bullet(x[1], x[0] + ': ')

# ===== 1. IPHONE =====
heading('1. IPHONE - cac van de nguoi dung quan tam & phan nan', 1)

heading('1.1. Hao pin & nong may sau cap nhat iOS (iOS 26.x)', 2); hot('cao')
bullet('Pin tut nhanh du chi nghe goi, nhan tin, luot Facebook/YouTube.')
bullet('May nong bat thuong ngay voi tac vu co ban; nong khi sac, choi game, dung 4G/5G.')
bullet('Apple da chinh thuc thua nhan iOS 26.1 gay hao pin, nong may.')
bullet('Nguyen nhan: tien trinh index/dong bo chay ngam sau update, app chua toi uu kip; thuong giam sau vai ngay.')

heading('1.2. iPhone 17 de tray xuoc - "Scratchgate"', 2); hot('cao')
bullet('Vien chuyen tu titan (15/16) sang nhom ren (17 Pro) - nhe & mat nhiet hon nhung mem, de xuoc hon.')
bullet('Lop anode quanh cum camera de troc, sut o canh sac; ban mau dam lo ro nhat.')
bullet('Tro thanh chu de viral tren Weibo/Facebook/X; Apple noi mot so vet la do de MagSafe cu, lau la sach.')

heading('1.3. Wi-Fi / Bluetooth khong on dinh', 2); hot('cao')
bullet('iPhone 17 bat song Wi-Fi yeu hon 15/16 o cung vi tri; Wi-Fi tu ngat roi ket noi lai.')
bullet('Bluetooth kho do thiet bi, tu ngat ket noi voi phu kien du song manh.')

heading('1.4. Camera & do ben thiet ke', 2); hot('tb')
bullet('Phan anh anh bi nhoe/mo tren iPhone 17 trong mot so truong hop.')
bullet('Mot so may Pro/Pro Max bi cong nhe, troc son sau thoi gian dung.')

heading('1.5. Loi phan cung hay gap (may da dung)', 2); hot('tb')
bullet('Face ID khong kha dung - do xung dot iOS, cum cam bien TrueDepth lech/ban bui, hoac da thay man hinh khong chuan.')
bullet('Man hinh: am vang, soc, diem chet - thuong do thay man kem chat luong.')
bullet('Loa re, mic nho, camera mo - do roi/vao nuoc hoac thay linh kien sai chuan.')
bullet('Luu y mua may cu: hu camera truoc co the keo theo mat Face ID (chung cum cap).')

# ===== 2. IPAD =====
heading('2. IPAD - cac van de pho bien', 1)
heading('2.1. Cam ung do, khong phan hoi', 2); hot('tb')
bullet('Chay nhieu app nang lam qua tai RAM, cam ung dung/do.')
heading('2.2. Sac khong vao pin', 2); hot('cao')
bullet('Day cap dut ngam/kem chat luong, chan sac ban hoac loi nguon; thu day chinh hang khac de loai tru.')
heading('2.3. Treo tao (Apple logo)', 2); hot('cao')
bullet('Man hinh dung o logo tao; thuong do xung dot phan mem, bo nho day, pin yeu/loi nguon.')
bullet('Co the kem sap nguon lien tuc, khong sac/khong ket noi may tinh - gay lo mat du lieu.')

# ===== 3. MACBOOK =====
heading('3. MACBOOK - cac van de pho bien', 1)
heading('3.1. Pin bi phong/phu', 2); hot('cao')
bullet('Dau hieu: ban phim & nap lung troi len, vo ho/nut, touchpad kho nhan, pin day nhung het nhanh.')
bullet('Nguy hiem: pin phong ep vao man hinh khi gap may, co the nut/vo man.')
heading('3.2. Sac khong vao pin', 2); hot('cao')
bullet('Do cuc/day sac hong, khong du dien; hoac pin chai, app chay ngam.')
heading('3.3. Cam sac van tut pin', 2); hot('tb')
bullet('Tai nang vuot cong suat sac, pin chai hoac loi mach sac.')
heading('3.4. He luy keo theo', 2); hot('thap')
bullet('Pin phong lam cong venh vo, ban phim, hong man hinh - chi phi sua/thay cao.')

# ===== 4. CHU DE XUYEN SUOT =====
heading('4. Nhung moi quan tam XUYEN SUOT moi dong san pham', 1)
bullet('Tuoi tho & do chai PIN - moi quan tam so 1 voi may da dung lau.', 'Pin: ')
bullet('Nong may khi choi game, quay video, sac - anh huong hieu nang & tuoi tho.', 'Nhiet: ')
bullet('Wi-Fi/Bluetooth/song - on dinh ket noi.', 'Ket noi: ')
bullet('Tray xuoc, cong, troc son - giu gia tri ban lai.', 'Do ben vo: ')
bullet('Chi phi sua cao, lo gap tho thay linh kien kem (mat Face ID, am man).', 'Sua chua & linh kien chinh hang: ')
bullet('Cap nhat iOS/macOS moi hay gay loi - nen update ngay hay cho?', 'Cap nhat phan mem: ')
bullet('iCloud day, chuyen du lieu sang may moi, mat du lieu khi treo tao.', 'Du lieu & iCloud: ')
bullet('Mua may dung/dinh iCloud/kich pin - noi lo lon nhat khi mua may cu.', 'Lua dao khi mua: ')

# ===== 5. GOI Y CONTENT =====
heading('5. Goi y khai thac content tu cac van de tren', 1)
para('Moi noi dau o tren = mot bai content. Cong thuc: goi dung loi -> giai thich nguyen nhan -> giai phap/dich vu cua shop.', italic=True, color=GREEN)
bullet('"iOS moi lam may nong & tut pin? Day la ly do va cach xu ly" - bai kien thuc hut tuong tac.')
bullet('"iPhone 17 de xuoc - 3 cach bao ve may giu gia ban lai" - bat trend Scratchgate.')
bullet('"Wi-Fi iPhone 17 chap chon? Thu 5 buoc nay" - giai dap loi dang hot.')
bullet('"Dau hieu pin MacBook phong nguy hiem - kiem tra ngay" - canh bao + dich vu thay pin.')
bullet('"iPad treo tao - cuu du lieu truoc khi qua muon" - dich vu cuu ho.')
bullet('"Mua may cu: 5 cach kiem tra pin, iCloud, Face ID" - tao niem tin, chot don.')

# ===== NGUON =====
heading('Nguon tham khao (cong khai)', 1)
sources = [
    'Apple thua nhan iOS 26.1 gay hao pin, nong may - Thanh Nien: https://thanhnien.vn/apple-thua-nhan-ios-261-gay-hao-pin-nong-may-185251106114403593.htm',
    'Loat loi tren iOS 26 - Dien Thoai Gia Kho: https://dienthoaigiakho.vn/tin-cong-nghe/loi-tren-ios-26/',
    'iOS 26.2 hao pin nhanh - Dienthoaivui: https://dienthoaivui.com.vn/tin-tuc/ios-26-2-hao-pin-nhanh',
    'iPhone 17 bi loi va cach khac phuc - XTmobile: https://www.xtmobile.vn/iphone-17-bi-loi',
    'iPhone 16 gap hang loat loi - Dan Tri: https://dantri.com.vn/cong-nghe/iphone-16-gap-hang-loat-loi-du-moi-ra-mat-20241025180917312.htm',
    'iPhone 17 Pro va iPhone Air mau toi de tray xuoc - VnExpress: https://vnexpress.net/iphone-17-pro-va-iphone-air-mau-toi-bi-phan-nan-de-tray-xuoc-4941792.html',
    'Apple ly giai vet xuoc tren iPhone 17 - Tuoi Tre: https://tuoitre.vn/apple-ly-giai-tin-don-xuat-hien-vet-xuoc-tren-cac-mau-iphone-17-20250927150610975.htm',
    '6 loi Face ID thuong gap - Thegioididong: https://www.thegioididong.com/hoi-dap/4-cach-khac-phuc-loi-faceid-tren-iphone-va-ipad-pr-1146094',
    'TOP 7 loi pin MacBook - Vien Di Dong: https://viendidong.com/loi-pin-macbook/',
    'MacBook sac khong vao pin - CellphoneS: https://cellphones.com.vn/sforum/macbook-sac-khong-vao-pin',
    'Loi iPad bi treo tao - CellphoneS: https://cellphones.com.vn/sforum/loi-ipad-bi-treo-tao',
    'Neu iPad khong sac - Apple Support: https://support.apple.com/en-us/102612',
]
for s in sources:
    p = doc.add_paragraph(style='List Bullet'); p.add_run(s).font.size = Pt(10)

out = '/home/user/DATA-MARKETING/Van_de_nguoi_dung_Apple_tren_MXH.docx'
doc.save(out)
print('Saved:', out)
